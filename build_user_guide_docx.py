"""Build Polly User Guide as a .docx formatted for KDP 6x9 trim."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup: 6x9 KDP with gutter ──
for section in doc.sections:
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.625)
    section.left_margin = Inches(0.75)  # gutter side
    section.right_margin = Inches(0.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Georgia'
font.size = Pt(11)
font.color.rgb = RGBColor(0x22, 0x22, 0x22)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = 'Georgia'
h1.font.size = Pt(18)
h1.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
h1.font.bold = True
h1.paragraph_format.space_before = Pt(0)
h1.paragraph_format.space_after = Pt(8)
h1.paragraph_format.page_break_before = True

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = 'Georgia'
h2.font.size = Pt(13)
h2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
h2.font.bold = True
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(4)

# Heading 3
h3 = doc.styles['Heading 3']
h3.font.name = 'Georgia'
h3.font.size = Pt(11)
h3.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
h3.font.bold = True
h3.paragraph_format.space_before = Pt(8)
h3.paragraph_format.space_after = Pt(3)

GREEN = RGBColor(0x05, 0x96, 0x69)
GRAY = RGBColor(0x88, 0x88, 0x88)
DARK = RGBColor(0x22, 0x22, 0x22)
AMBER = RGBColor(0x92, 0x40, 0x0E)

def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None, space_before=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    return p

def add_mixed(parts, align=None, space_after=None):
    """parts = list of (text, bold, italic, color, size) tuples"""
    p = doc.add_paragraph()
    for text, bold, italic, color, size in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if color: run.font.color.rgb = color
        if size: run.font.size = Pt(size)
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_voice(say_text, reply_text=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(say_text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)
    run.font.size = Pt(10)
    if reply_text:
        run2 = p.add_run('\n' + reply_text)
        run2.italic = True
        run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        run2.font.size = Pt(10)

def add_tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run('Tip: ')
    run.bold = True
    run.font.color.rgb = AMBER
    run.font.size = Pt(9.5)
    run2 = p.add_run(text)
    run2.font.size = Pt(9.5)

def add_feature(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x05, 0x50, 0x50)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)

def add_numbered(text, bold_parts=None):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════
#  COVER PAGE
# ════════════════════════════════════════════
# Add some spacing to center vertically
for _ in range(6):
    add_para('', space_after=0)

# Try to add logo
logo_path = 'C:/Users/big_g/Desktop/polly-connect/polly_logo_from_doc.png'
if os.path.exists(logo_path):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(2.4))

add_para('User Guide', bold=True, size=26, color=GREEN, align='center', space_before=8)
add_para('Your Friendly Talking Companion', italic=True, size=13, color=GRAY, align='center', space_after=24)

add_para('A complete guide for owners, caretakers, and families.', size=11, color=RGBColor(0x44,0x44,0x44), align='center', space_after=2)
add_para('Everything you need to set up Polly, capture stories,', size=11, color=RGBColor(0x44,0x44,0x44), align='center', space_after=2)
add_para('and build a Legacy Book for generations to come.', size=11, color=RGBColor(0x44,0x44,0x44), align='center', space_after=36)

add_mixed([
    ('polly-connect.com', True, False, GREEN, 11),
], align='center', space_after=2)
add_para('Capturing memories, one story at a time.', italic=True, size=9, color=GRAY, align='center')

page_break()

# ════════════════════════════════════════════
#  QUICK START PAGE
# ════════════════════════════════════════════
add_para('🦜', size=28, align='center', space_after=0)
add_para('Polly Connect', bold=True, size=20, color=GREEN, align='center', space_after=0)
add_para('Capturing memories, one story at a time', italic=True, size=9, color=GRAY, align='center', space_after=12)

add_para('— QUICK START —', bold=True, size=12, color=GREEN, align='center', space_after=8)

add_para('YOUR CLAIM CODE', bold=True, size=9, color=AMBER, align='center', space_after=4)
add_para('___  ___  ___  ___  ___  ___', bold=True, size=22, color=AMBER, align='center', space_after=12)

add_numbered('Go to polly-connect.com and click Register.')
add_numbered('Plug in Polly and connect to "Polly-Setup" WiFi on your phone.')
add_numbered('Enter the 6-digit claim code above and your home WiFi password, tap Connect.')
add_numbered('On the website: Settings → Devices → enter the same claim code.')
add_numbered('Done! Say "Hey Polly" and start sharing memories.')

add_para('')
add_para('Every family has a story worth keeping.', bold=True, size=10, color=GREEN, align='center', space_after=2)
add_para('Polly listens, remembers, and reads them back with love.', italic=True, size=8, color=GRAY, align='center', space_after=16)

# Family code section
add_para('─' * 50, color=GRAY, align='center', space_after=4)
add_para('FAMILY ACCESS CODE', bold=True, size=10, color=GREEN, align='center', space_after=4)
add_para('___  ___  ___  ___  ___  ___', bold=True, size=18, color=GREEN, align='center', space_after=4)
add_para('Share with family so they can view stories, photos & the Legacy Book.', italic=True, size=8, color=GRAY, align='center', space_after=2)
add_para('Generate from Settings on the web portal.', italic=True, size=8, color=GRAY, align='center')

page_break()

# ════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════
add_para('Contents', bold=True, size=18, color=GREEN, align='center', space_after=16)

toc_entries = [
    ('Getting Started', '1–2', 'What Is Polly? · Setting Up · Welcome Page'),
    ('Talking to Polly', '3–5', 'Jokes · Bible · Weather · Meds · Items · Messages'),
    ('Stories & Memories', '6–7', 'Story Sessions · Record Button · Phone Recording · Transcriptions'),
    ('The Legacy Book', '8', 'Chapters · Customizable Narratives · Exporting & Printing'),
    ('Photos & QR Codes', '9', 'Uploading · Photo Stories · Book Toggle · QR Codes'),
    ('Nostalgia & Prayers', '10–11', 'Nostalgia Snippets · Learning System · AI Prayers · Prayer Requests'),
    ('Family & Sharing', '12', 'Access Codes · Family Tree · Permissions'),
    ('The Web Portal', '13–14', 'Dashboard · Settings · Pronunciation Guide'),
    ('Voice Command Reference', '15', None),
    ('Tips & Troubleshooting', '16', None),
    ('Notes', '17', None),
]

for title, pages, subtitle in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    # Add dots and page
    run2 = p.add_run('  ' + '·' * 30 + '  ' + pages)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = GRAY
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Inches(0.15)
        run3 = p2.add_run(subtitle)
        run3.font.size = Pt(9.5)
        run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

page_break()

# ════════════════════════════════════════════
#  GETTING STARTED
# ════════════════════════════════════════════
doc.add_heading('Getting Started', level=1)

doc.add_heading('What Is Polly?', level=2)
doc.add_paragraph('Polly is a friendly talking companion designed for seniors and families. She sits on a table or countertop and listens for your voice. Just say "Hey Polly" to wake her up, then ask her anything!')
doc.add_paragraph('Polly can tell jokes, read Bible verses, share the weather, remember where you put things, help you keep track of medications, pray with you, take you on a trip down memory lane, and — most importantly — help you tell your life story for future generations.')
doc.add_paragraph('She connects to your home WiFi and works through a small speaker and microphone. No screens to read, no buttons to press — just talk to her like a friend.')

doc.add_heading('Who Is This Guide For?', level=2)
p = doc.add_paragraph()
run = p.add_run('This guide is for the ')
run = p.add_run('caretaker')
run.bold = True
p.add_run(' — the family member who sets Polly up and manages her settings through the web portal. The person Polly talks to (the "owner") doesn\'t need to read this guide. They just need to know the wake word: ')
run = p.add_run('"Hey Polly."')
run.bold = True
doc.add_paragraph('Throughout this guide, we\'ll use "Grandma" as the owner and "you" as the caretaker, but Polly works for anyone.')

add_para('· · ·', align='center', color=GREEN, size=12, space_before=8, space_after=8)

doc.add_heading('Setting Up Polly', level=2)
doc.add_heading('Step 1: Plug In & Connect to WiFi', level=3)
doc.add_paragraph('Plug Polly into a power outlet using the included USB cable. When she powers on, her light will blink slowly — she\'s looking for WiFi.')
doc.add_paragraph('If Polly can\'t find a saved network, she creates a temporary one called "Polly-Setup":')
add_numbered('On your phone, go to WiFi settings.')
add_numbered('Connect to "Polly-Setup".')
add_numbered('A setup page pops up automatically. (If not, go to 192.168.4.1 in your browser.)')
add_numbered('Select your home WiFi and enter the password.')
add_numbered('Enter your 6-digit claim code from the Quick Start card.')
add_numbered('Tap Connect. Polly reboots and connects. The light stops blinking when she\'s ready.')

add_tip('Need to change WiFi? Hold the small button for 3 seconds while plugging in. This resets WiFi and puts Polly back in setup mode.')

doc.add_heading('Step 2: Create Your Account', level=3)
p = doc.add_paragraph('Go to ')
run = p.add_run('polly-connect.com')
run.bold = True
run.font.color.rgb = GREEN
p.add_run(' and click ')
run = p.add_run('Register')
run.bold = True
p.add_run('. Create an account with your email and password.')

doc.add_heading('Step 3: Claim Your Device', level=3)
p = doc.add_paragraph('Go to ')
run = p.add_run('Settings → Devices')
run.bold = True
p.add_run(' and enter the same ')
run = p.add_run('6-digit claim code')
run.bold = True
p.add_run('. This links Polly to your account.')

page_break()

doc.add_heading('Your First-Time Welcome Page', level=2)
# Remove the page break that heading 2 doesn't add
doc.add_paragraph('After creating your account, Polly walks you through a Welcome page to get started:')
add_bullet('who Polly talks to (e.g., "Grandma Helen")', bold_prefix='Owner\'s name — ')
add_bullet('what Polly calls them (e.g., "Helen" or "Mom")', bold_prefix='Familiar name — ')
add_bullet('where they grew up (used for nostalgia and weather)', bold_prefix='Hometown — ')
add_bullet('helps Polly generate era-appropriate memories', bold_prefix='Birth year — ')
add_bullet('for accurate weather forecasts', bold_prefix='Location — ')

add_tip('You can always change these later in Settings. The Welcome page just gets the essentials filled in so Polly can start personalizing right away.')

add_para('· · ·', align='center', color=GREEN, size=12, space_before=16, space_after=8)
add_para('3 Easy Steps', bold=True, size=14, color=GREEN, align='center', space_after=4)
add_para('That\'s it — Polly is ready to go. The owner just needs to know one thing:', size=10, color=GRAY, align='center', space_after=8)
add_para('Say "Hey Polly"', bold=True, size=16, color=GREEN, align='center', space_after=8)
add_para('Everything else — stories, settings, the Legacy Book — is managed by you through the web portal at polly-connect.com.', size=10, color=GRAY, align='center')

# ════════════════════════════════════════════
#  TALKING TO POLLY
# ════════════════════════════════════════════
doc.add_heading('Talking to Polly', level=1)
doc.add_paragraph('Every conversation starts with the wake word:')
add_voice('Say: "Hey Polly"', 'Polly: "Well hello there! What can I do for you?"')
doc.add_paragraph('After Polly responds, she goes back to sleep. Say "Hey Polly" again for your next question. There\'s no rush — she\'s very patient.')

doc.add_heading('Greetings & Goodbyes', level=2)
add_voice('Say: "Good morning"', 'Polly: "Good morning! It\'s wonderful to hear your voice today."')
add_voice('Say: "Goodbye"', 'Polly: "Goodbye! It was lovely talking with you."')
add_voice('Say: "Thank you"', 'Polly: "You\'re very welcome, dear."')

doc.add_heading('Jokes', level=2)
doc.add_paragraph('Polly knows over a thousand jokes and pauses before the punchline for dramatic effect!')
add_voice('Say: "Tell me a joke"')
add_voice('Say: "Tell me another joke"')
doc.add_paragraph('For the grandkids — 100 kid-friendly jokes:')
add_voice('Say: "Tell me a kid joke"')
add_voice('Say: "Tell me a fart joke" · "Tell me a dinosaur joke"')
add_tip('Polly never repeats the same joke twice in a row!')

doc.add_heading('Bible Verses', level=2)
doc.add_paragraph('Polly knows 336 Bible verses and can read them by topic or book.')
add_voice('Say: "Read me a Bible verse"')
add_voice('Say: "Read me a verse about hope"')
add_voice('Say: "Read me a Psalm"', 'Polly: "Psalm 23:4 — Even though I walk through the darkest valley, I will fear no evil, for you are with me."')

doc.add_heading('Weather', level=2)
doc.add_paragraph('Polly gets the real weather for your location automatically.')
add_voice('Say: "What\'s the weather?"', 'Polly: "Right now in Columbus, Ohio, it\'s 72 degrees and partly cloudy."')
add_voice('Say: "Do I need an umbrella?"')
doc.add_paragraph('Polly also shares a fun Farmer\'s Almanac fact with every weather report!')

doc.add_heading('Time & Date', level=2)
add_voice('Say: "What time is it?"', 'Polly: "It\'s 2:30 in the afternoon."')
add_voice('Say: "What day is it?"', 'Polly: "Today is Saturday, March 8th, 2026."')

doc.add_heading('Medications', level=2)
doc.add_paragraph('Set up medications in the web portal (name, dosage, times). Polly reminds the owner and tracks what they\'ve taken.')
add_voice('Say: "What are my medications?"', 'Polly: "You have 2 medications: Metformin at 8 AM and 6 PM, and Lisinopril at 8 AM."')
add_voice('Say: "I took my medicine"', 'Polly: "Great job! I\'ve logged that for today."')
add_tip('View medication logs and export a calendar (ICS file) from Care → Medications on the web portal.')

doc.add_heading('Storing & Finding Items', level=2)
doc.add_paragraph('Polly remembers where you put things — glasses, keys, remotes, all those things that walk away on their own!')
add_voice('Say: "My keys are on the counter"', 'Polly: "Got it. The keys are on the counter."')
add_voice('Say: "Where are my keys?"', 'Polly: "The keys are on the counter."')
add_voice('Say: "What\'s on the counter?"', 'Polly: "On the counter, you have your keys and your glasses."')
add_voice('Say: "Forget about the keys"', 'Polly: "Forgot about the keys."')

doc.add_heading('Family Message Board', level=2)
doc.add_paragraph('A voice-powered bulletin board for the whole family.')
doc.add_heading('Leaving Messages', level=3)
add_voice('Say: "Tell Dad I\'m going to the store"', 'Polly: "Got it. I\'ll let Dad know."')
doc.add_heading('Checking Messages', level=3)
add_voice('Say: "Any messages?"', 'Polly: "You have 2 messages. Glen says: I love you so much."')
add_voice('Say: "Where is Dad?"', 'Polly: "Dad is going to work."')
doc.add_heading('Clearing Messages', level=3)
add_voice('Say: "Dad is home"', 'Polly: "Welcome home! I\'ll clear Dad\'s messages."')
add_tip('Messages expire after 24 hours automatically. You can also send messages from the web portal.')

doc.add_heading('Family Lookup', level=2)
doc.add_paragraph('If you\'ve added family to the Family Tree on the web portal, Polly can tell you about them.')
add_voice('Say: "Who is Mia?"', 'Polly: "Mia is your granddaughter."')

doc.add_heading('Keeping Polly Quiet', level=2)
doc.add_paragraph('If Polly is squawking or chattering and you need some peace:')
add_voice('Say: "Be quiet" or "Hush"', 'Polly: "Okay okay! I\'ll be quiet."')
add_tip('She\'ll go quiet for a while, but she\'ll start chattering again eventually — she is a parrot, after all! Adjust squawk volume and quiet hours in Settings.')

# ════════════════════════════════════════════
#  STORIES & MEMORIES
# ════════════════════════════════════════════
doc.add_heading('Stories & Memories', level=1)
doc.add_paragraph('This is the heart of Polly. She asks thoughtful questions about life and records the answers. Over time, these stories become a Legacy Book — a printed book of family memories that lasts for generations.')

doc.add_heading('Family Story Sessions', level=2)
add_voice('Say: "Ask me a family question"', 'Polly: "Here\'s one for you: What\'s your earliest childhood memory?"')
doc.add_paragraph('Just start talking! Polly listens patiently, even if you pause to think. When you finish, she might ask a follow-up to dig deeper.')
add_voice('Say: "I\'m done" — end the story')
add_voice('Say: "Skip" — get a different question')
add_voice('Say: "How many stories do I have?" — check progress')

doc.add_heading('Ask About Specific Topics', level=3)
add_voice('Say: "Tell me a story about fishing"', 'Polly shares a relevant story or asks about that topic if none exist yet.')

add_tip('Don\'t worry about saying everything perfectly. Polly captures your words and you can review and correct them later. Stories are fully customizable — edit text, fix names, add details anytime.')

doc.add_heading('Story Record Button', level=2)
doc.add_paragraph('Polly has a physical button (K1/+ on the side) for recording longer stories without a wake word.')
add_voice('Press the K1/+ button once:', 'Polly: "Recording started. Take your time and tell me your story."')
doc.add_paragraph('The LED turns solid while recording. Just talk naturally. Press the button again to stop. Recording auto-stops after 30 minutes.')
add_tip('Perfect for long stories, family gatherings, or when Grandma just wants to talk without saying "Hey Polly" between thoughts.')

doc.add_heading('Record a Memory from Your Phone', level=2)
doc.add_paragraph('You don\'t need to be in the same room as Polly to capture memories. The web portal has a "Record a Memory" button on the Stories page.')
add_numbered('Go to Stories on the web portal.')
add_numbered('Tap "Record a Memory" at the top.')
add_numbered('Enter the speaker\'s name.')
add_numbered('Tap the red record button and start talking. You\'ll see a volume indicator and timer.')
add_numbered('Tap Stop when finished.')

add_feature('Audio memories are always saved — even if Polly can\'t transcribe the words (kids playing, laughter at a gathering, ambient sounds), the audio recording is kept. These sounds become part of your family\'s memory collection.')

doc.add_heading('Fully Customizable Transcriptions', level=2)
doc.add_paragraph('Every story Polly transcribes can be reviewed and corrected on the web portal:')
add_bullet('fix names, fill in details, correct anything Polly misheard.', bold_prefix='Edit text — ')
add_bullet('mark a transcription as reviewed and accurate.', bold_prefix='Verify — ')
add_bullet('Polly auto-tags names, but you can add more.', bold_prefix='Tag people & places — ')
add_bullet('attach an approximate year to help place the story in the right chapter.', bold_prefix='Add dates — ')
add_tip('Family members with an access code can also verify transcriptions from their own phone!')

# ════════════════════════════════════════════
#  THE LEGACY BOOK
# ════════════════════════════════════════════
doc.add_heading('The Legacy Book', level=1)
doc.add_paragraph('This is what makes Polly truly special. Over time, Polly takes all of the owner\'s stories and weaves them into a beautiful Legacy Book — organized by life chapters, enriched with photos and audio QR codes, and ready to print.')

doc.add_heading('How the Book Works', level=2)
doc.add_paragraph('Polly organizes stories into 14 chapters covering different themes of life — Childhood & Early Years, Family & Home, Love & Relationships, Career & Accomplishments, Adventures & Travel, Faith & Values, and more.')
doc.add_paragraph('When enough stories are collected for a chapter, AI writes a chapter draft that weaves memories into a warm, flowing narrative. Each chapter builds on the last — the AI remembers what it\'s written and avoids repeating itself.')

doc.add_heading('Fully Customizable Narratives', level=2)
doc.add_paragraph('You have complete control over every chapter:')
add_bullet('rewrite sentences, add details, fix facts.', bold_prefix='Edit the narrative — ')
add_bullet('love it? Lock it in. Not right? Regenerate with updated stories.', bold_prefix='Keep or regenerate — ')
add_bullet('ask Polly to read any chapter aloud.', bold_prefix='Replay by voice — ')
add_bullet('Polly uses birth years, story dates, and family milestones to place events in the right era.', bold_prefix='Timeline-enriched — ')

add_feature('The more stories, the richer the book. Every story, photo, and correction makes the next chapter draft better. There\'s no limit — keep adding memories and the book keeps growing.')

doc.add_heading('Exporting & Printing Your Book', level=2)
doc.add_paragraph('Go to Stories → Book Export on the web portal. Polly generates a beautifully formatted PDF with:')
add_bullet('Chapter narratives woven from real stories')
add_bullet('Family photos placed alongside relevant chapters')
add_bullet('QR codes that link to original audio recordings')
add_bullet('A professional layout ready for printing or binding')
add_tip('Order a printed copy through any print-on-demand service, or print at home. The PDF is formatted and ready to go.')

# ════════════════════════════════════════════
#  PHOTOS & QR CODES
# ════════════════════════════════════════════
doc.add_heading('Photos & QR Codes', level=1)
doc.add_paragraph('The Photos section is where family memories come alive. Upload old photos, record stories about them, and control which ones appear in the Legacy Book.')

doc.add_heading('Uploading & Tagging Photos', level=2)
add_numbered('Go to Photos on the web portal.')
add_numbered('Tap Upload Photo.')
add_numbered('Choose a photo from your phone or computer.')
add_numbered('Add a caption ("Christmas 1985"), date, and tags (names of people).')
add_numbered('Tap Upload.')
add_tip('Tag photos with family member names! Tap a person on the Family Tree to see all their tagged photos.')

doc.add_heading('Recording Photo Stories', level=2)
doc.add_paragraph('Every photo has a microphone button — one of Polly\'s most powerful features:')
add_numbered('Tap the microphone button on any photo.')
add_numbered('Enter your name.')
add_numbered('Tap Record — talk about who\'s in the photo, when it was taken, what was happening.')
add_numbered('Tap Stop when done.')
doc.add_paragraph('Photo stories feed directly into the Legacy Book. The AI chapter writer references photos naturally in the narrative.')
add_tip('Family members with an access code can upload photos and record stories from anywhere in the world!')

doc.add_heading('Fully Customizable Book Inclusion', level=2)
doc.add_paragraph('Every photo has a sliding toggle that controls whether it appears in the Legacy Book:')
add_bullet('(green) — included in the book PDF.', bold_prefix='Toggle ON ')
add_bullet('(gray) — kept in the gallery but not printed.', bold_prefix='Toggle OFF ')
doc.add_paragraph('The toggle appears on all photos — with or without a linked story. Polly confirms before removing a photo from the book.')

doc.add_heading('QR Codes — Hear the Story', level=2)
doc.add_paragraph('Stories with audio get a QR code in the Legacy Book. Scan it with your phone and hear the original voice telling the story.')
add_feature('Imagine this: Your grandchildren open the Legacy Book fifty years from now, scan a QR code, and hear Great-Grandma\'s voice telling the story of how she met Great-Grandpa. That\'s the power of Polly.')

# ════════════════════════════════════════════
#  NOSTALGIA & PRAYERS
# ════════════════════════════════════════════
doc.add_heading('Nostalgia & Prayers', level=1)

doc.add_heading('Nostalgia Snippets', level=2)
doc.add_paragraph('Polly can take the owner on a trip down memory lane! Based on their hometown and birth year, Polly generates era-appropriate memories — drive-in theaters, local diners, the cars they drove, the music they listened to.')
add_voice('Say: "Take me back"', 'Polly: "Remember cruising down Main Street on Friday nights? The A&W was packed, the jukebox was playing Elvis, and everybody knew everybody…"')
doc.add_paragraph('Polly also sprinkles nostalgia into her regular chatter — about 20% of the time she talks on her own, she\'ll share a memory.')

doc.add_heading('Setting Up Nostalgia', level=3)
add_numbered('Go to Family → Nostalgia on the web portal.')
add_numbered('Fill in the Nostalgia Profile — what kind of kid they were, sports, hangouts, cars, first job, and more.')
add_numbered('Click Save Profile, then Generate Snippets.')
add_numbered('Polly creates 25 personalized memories across categories: Hometown, Sports, Cars, Music, Food, Culture, Childhood, Military, and Work.')

doc.add_heading('Fully Customizable Nostalgia Learning', level=2)
doc.add_paragraph('The nostalgia system learns from your feedback:')
add_bullet('Polly treats your correction as fact. Change "Coach Peterson" to "Coach Zuke" and Polly remembers Coach Zuke forever. Edited snippets show a blue "Edited" badge.', bold_prefix='Edit a snippet — ')
add_bullet('tells Polly "more like this." She generates similar themes next time.', bold_prefix='Keep a snippet — ')
add_bullet('Polly logs it and avoids those topics in the future.', bold_prefix='Delete a snippet — ')
add_bullet('adds new snippets without replacing the ones you\'ve curated. Your edits, keeps, and deletes all inform the next batch.', bold_prefix='Generate More — ')
add_feature('The nostalgia collection grows and improves over time. The more you curate it, the more personal and accurate it becomes. Polly\'s AI learns from every edit, keep, and delete.')

doc.add_heading('AI Prayers', level=2)
doc.add_paragraph('Polly can pray with the owner using AI-generated prayers that feel personal and meaningful.')
add_voice('Say: "Say a prayer"')
add_voice('Say: "Pray for peace"')
add_voice('Say: "I\'m having a hard day"')
add_voice('Say: "Pray for Mia"')
doc.add_paragraph('Polly draws from 15 emotional categories — hope, anxiety, grief, loneliness, gratitude, strength, healing, family, peace, faith, forgiveness, purpose, joy, rest, and guidance. She rotates so prayers feel fresh.')
doc.add_paragraph('Prayers are family-aware: Polly mentions close family by name. She\'s also deceased-aware — loved ones who have passed are spoken of with love and gratitude.')

doc.add_heading('Prayer Requests', level=2)
doc.add_paragraph('Add specific prayer requests through the web portal:')
add_numbered('Go to Family → Prayers.')
add_numbered('Add names and requests (e.g., "Pray for John\'s recovery").')
add_numbered('Polly weaves them naturally into her prayers.')
add_numbered('Remove requests when they\'re no longer needed.')

# ════════════════════════════════════════════
#  FAMILY & SHARING
# ════════════════════════════════════════════
doc.add_heading('Family & Sharing', level=1)

doc.add_heading('Family Access Codes', level=2)
doc.add_paragraph('Want family to see Grandma\'s stories, photos, and Legacy Book? Generate a 6-digit family code:')
add_numbered('Go to Settings on the web portal.')
add_numbered('Under Family Access Code, tap Generate Code.')
add_numbered('Share the code with family.')
add_numbered('They go to polly-connect.com/web/family, enter the code and their name.')
add_tip('Regenerate the code anytime to change who has access. The old code stops working immediately. Write your code on the Quick Start card at the front of this guide!')

doc.add_heading('Family Tree', level=2)
doc.add_paragraph('Add family members with their relationship to the owner. The visual tree shows up to 8 generations. Each member can have:')
add_bullet('Name, relationship, and spouse name')
add_bullet('Birth year and a short bio')
add_bullet('Deceased flag (Polly speaks of them with love and respect)')
doc.add_paragraph('Tap any person to see all their tagged photos.')

doc.add_heading('What Family Members Can Do', level=2)
add_bullet('View all stories, transcriptions, and the Legacy Book')
add_bullet('Browse and upload family photos')
add_bullet('Record photo stories from their phone')
add_bullet('View the family tree and add new members')
add_bullet('Send messages on the family message board')
add_bullet('Verify transcriptions for accuracy')
p = doc.add_paragraph('They ')
run = p.add_run('cannot')
run.bold = True
p.add_run(' change settings, delete content, manage medications, or access prayers and nostalgia — that\'s the caretaker\'s job.')

# ════════════════════════════════════════════
#  THE WEB PORTAL
# ════════════════════════════════════════════
doc.add_heading('The Web Portal', level=1)
doc.add_paragraph('Everything Polly captures can be reviewed and managed at polly-connect.com. The portal is organized into clear sections:')

doc.add_heading('Dashboard', level=2)
doc.add_paragraph('Your home base. Story count, medication alerts, Legacy Book progress, and recent activity at a glance. Big buttons take you to Stories, Family, Care, and Settings.')

doc.add_heading('Stories', level=2)
add_bullet('all recorded stories with search and edit', bold_prefix='Stories — ')
add_bullet('review and verify what Polly heard', bold_prefix='Transcriptions — ')
add_bullet('Legacy Book chapters and progress', bold_prefix='Book — ')
add_bullet('download the printable PDF', bold_prefix='Book Export — ')

doc.add_heading('Family', level=2)
add_bullet('visual tree up to 8 generations', bold_prefix='Family Tree — ')
add_bullet('gallery with upload, tags, and book toggles', bold_prefix='Photos — ')
add_bullet('family message board', bold_prefix='Messages — ')
add_bullet('stored items and locations', bold_prefix='Memory — ')
add_bullet('manage prayer requests (owner only)', bold_prefix='Prayers — ')
add_bullet('manage nostalgia snippets (owner only)', bold_prefix='Nostalgia — ')

doc.add_heading('Care', level=2)
add_bullet('add, edit, and track medications', bold_prefix='Medications — ')
add_bullet('export schedule as ICS file', bold_prefix='Med Calendar — ')

doc.add_heading('Settings & Preferences', level=2)
doc.add_paragraph('Settings are organized into collapsible sections:')

doc.add_heading('User Profile', level=3)
doc.add_paragraph('Name, familiar name, hometown, birth year, location, and language preferences.')
doc.add_heading('Preferences', level=3)
doc.add_paragraph('Memory care mode, quiet hours, squawk intervals, and behavior settings.')
doc.add_heading('Polly Sounds', level=3)
doc.add_paragraph('Adjust squawk volume (0–100%) and voice volume. Find the right level — loud enough to hear, quiet enough not to startle.')
doc.add_heading('Pronunciation & Voice Sensitivity', level=3)
doc.add_paragraph('If Polly mispronounces a name, add it to the Pronunciation Guide. Type the name and how it should sound (e.g., "Liam" → "LEE-um"). The fix applies everywhere — stories, prayers, narratives, and chatter.')
doc.add_paragraph('Adjust microphone sensitivity if Polly picks up too much background noise or isn\'t hearing well.')
doc.add_heading('Snooze Polly', level=3)
doc.add_paragraph('Temporarily silence squawks and chatter. Great for nap time or visitors.')
doc.add_heading('Family Access Code', level=3)
doc.add_paragraph('Generate, view, or regenerate your 6-digit sharing code.')
doc.add_heading('Devices', level=3)
doc.add_paragraph('View connected devices, claim new ones, print setup cards, and manage device settings.')

# ════════════════════════════════════════════
#  VOICE COMMAND REFERENCE
# ════════════════════════════════════════════
doc.add_heading('Voice Command Reference', level=1)

commands = [
    ('"Hey Polly"', 'Wake up'),
    ('"Tell me a joke"', 'Random joke'),
    ('"Tell me a kid joke"', 'Kid-friendly joke'),
    ('"Read me a Bible verse"', 'Random verse'),
    ('"Read me a verse about hope"', 'Verse by topic'),
    ('"What\'s the weather?"', 'Current forecast'),
    ('"What time is it?"', 'Current time'),
    ('"What day is it?"', 'Current date'),
    ('"What are my medications?"', 'Med list & schedule'),
    ('"I took my medicine"', 'Log medication'),
    ('"Ask me a family question"', 'Start story session'),
    ('"How many stories do I have?"', 'Story count'),
    ('"Tell me a story about…"', 'Topic-specific story'),
    ('"My keys are on the counter"', 'Store item location'),
    ('"Where are my keys?"', 'Find item'),
    ('"What\'s on the counter?"', 'Items at location'),
    ('"Forget about the keys"', 'Remove item'),
    ('"Any messages?"', 'Check message board'),
    ('"Tell Dad I love him"', 'Leave a message'),
    ('"Where is Dad?"', 'Check person status'),
    ('"Dad is home"', 'Clear person\'s messages'),
    ('"Clear the board"', 'Clear all messages'),
    ('"Say a prayer"', 'AI prayer'),
    ('"Pray for Mia"', 'Person-specific prayer'),
    ('"I\'m having a hard day"', 'Comforting prayer'),
    ('"Take me back"', 'Nostalgia snippet'),
    ('"Who is Mia?"', 'Family lookup'),
    ('"Be quiet" / "Hush"', 'Silence squawks'),
    ('"Skip"', 'Next question'),
    ('"I\'m done"', 'End story session'),
    ('"Repeat that"', 'Replay last response'),
    ('"Slower"', 'Replay slower'),
    ('"Goodbye"', 'End conversation'),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
# Header row
hdr = table.rows[0].cells
for i, text in enumerate(['What to Say', 'What Polly Does']):
    hdr[i].text = text
    run = hdr[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREEN

for cmd, action in commands:
    row = table.add_row().cells
    row[0].text = cmd
    row[1].text = action
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.space_before = Pt(1)

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(2.3)
    row.cells[1].width = Inches(2.2)

# ════════════════════════════════════════════
#  TIPS & TROUBLESHOOTING
# ════════════════════════════════════════════
doc.add_heading('Tips & Troubleshooting', level=1)

doc.add_heading('Speaking Tips', level=2)
add_bullet('Speak naturally — Polly understands conversational language.')
add_bullet('Pause after "Hey Polly" and wait for her greeting before speaking.')
add_bullet('If Polly says "I didn\'t understand that," try rephrasing.')
add_bullet('Don\'t rush. Polly is designed for patient, natural conversation.')
add_bullet('Reduce background noise (TV, radio) for best results.')

doc.add_heading('Common Issues & Fixes', level=2)

doc.add_heading('Polly isn\'t responding', level=3)
add_bullet('Make sure she\'s plugged in and the light is on.')
add_bullet('Check that your WiFi is working.')
add_bullet('Try unplugging and plugging back in.')

doc.add_heading('Polly says "I didn\'t understand" too often', level=3)
add_bullet('Say "Hey Polly" clearly before each command.')
add_bullet('Reduce background noise.')
add_bullet('Speak at normal volume.')
add_bullet('Adjust microphone sensitivity in Settings.')

doc.add_heading('Polly needs new WiFi', level=3)
add_bullet('Hold the small button for 3 seconds while plugging in.')
add_bullet('Connect to "Polly-Setup" and enter new credentials.')

doc.add_heading('Polly mispronounces a name', level=3)
add_bullet('Go to Settings → Pronunciation & Voice Sensitivity.')
add_bullet('Add the name and how Polly should say it.')

doc.add_heading('Polly is too loud or too quiet', level=3)
add_bullet('Settings → Polly Sounds → adjust volume sliders.')

# ════════════════════════════════════════════
#  NOTES
# ════════════════════════════════════════════
doc.add_heading('Notes', level=1)
add_para('Use this space for family contacts, story ideas, or anything else.', size=9, color=GRAY, space_after=8)

add_para('Family Members & Contact Info', bold=True, size=9, color=GRAY, space_after=2)
for _ in range(8):
    add_para('_' * 60, size=10, color=RGBColor(0xDD, 0xDD, 0xDD), space_after=6)

add_para('Story Ideas & Questions to Ask', bold=True, size=9, color=GRAY, space_after=2, space_before=8)
for _ in range(8):
    add_para('_' * 60, size=10, color=RGBColor(0xDD, 0xDD, 0xDD), space_after=6)

add_para('Other Notes', bold=True, size=9, color=GRAY, space_after=2, space_before=8)
for _ in range(6):
    add_para('_' * 60, size=10, color=RGBColor(0xDD, 0xDD, 0xDD), space_after=6)

page_break()

# ════════════════════════════════════════════
#  BACK COVER
# ════════════════════════════════════════════
for _ in range(5):
    add_para('', space_after=0)

add_para('"Every family has a story worth keeping."', bold=False, italic=True, size=15, color=GREEN, align='center', space_after=4)
add_para('Polly listens, remembers, and reads them back with love.', italic=True, size=10, color=GRAY, align='center', space_after=24)

features = [
    'Voice-activated — no screens, no buttons, just conversation',
    'Fully customizable stories, photos, and narratives',
    'AI-powered Legacy Book with printable PDF export',
    'QR codes that play the original voice recording',
    'Personalized nostalgia that learns and improves',
    'AI prayers with 15 emotional categories',
    'Family sharing — everyone can contribute memories',
    'Phone recording — capture memories from anywhere',
    '1,000+ jokes, 336 Bible verses, weather, medications',
    'Item memory — "Where are my keys?"',
    'Family message board — voice-powered family notes',
    'Pronunciation guide — every name said right',
]
for f in features:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run('✓  ')
    run.font.color.rgb = GREEN
    run.bold = True
    run.font.size = Pt(9.5)
    run2 = p.add_run(f)
    run2.font.size = Pt(9.5)

add_para('', space_after=16)
add_para('polly-connect.com', bold=True, size=11, color=GREEN, align='center', space_after=4)
add_para('Patent Pending', size=8, color=GRAY, align='center')

# ── Add page numbers via footer ──
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add auto page number field
    run = p.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    fld_xml = (
        '<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    run._r.append(parse_xml(fld_xml))
    run2 = p.add_run(' PAGE ')
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY
    fld_instr = parse_xml(
        '<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>'
    )
    run2._r.clear()
    run2._r.append(fld_instr)
    run3 = p.add_run()
    run3.font.size = Pt(9)
    fld_sep = parse_xml(
        '<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    run3._r.append(fld_sep)
    run4 = p.add_run('1')
    run4.font.size = Pt(9)
    run4.font.color.rgb = GRAY
    run5 = p.add_run()
    run5.font.size = Pt(9)
    fld_end = parse_xml(
        '<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
    )
    run5._r.append(fld_end)

# ── Save ──
output_path = 'C:/Users/big_g/Desktop/Polly User Guide v2 updated.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
print('Done!')
